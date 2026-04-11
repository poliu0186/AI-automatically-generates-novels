import json
from datetime import datetime, timedelta
from decimal import Decimal, InvalidOperation
from io import BytesIO

from flask import Blueprint, current_app, jsonify, request, send_file
from flask_login import current_user, login_required

from app.activity_logging import log_user_action
from app.billing import InsufficientBalanceError, apply_recharge, charge_feature_coins, create_recharge_order, get_or_create_wallet
from app.extensions import db
from app.models import AdminSetting, ExportedArticle, LLMUsageRecord, RechargeOrder, WalletLedger

try:
    from docx import Document
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate
    PDF_AVAILABLE = True
except Exception:
    PDF_AVAILABLE = False

payment_bp = Blueprint('payment', __name__)


def _get_admin_setting(key, default=''):
    row = AdminSetting.query.filter_by(key=key).first()
    if row and row.value is not None:
        return row.value
    return default


def _payment_channel():
    return (_get_admin_setting('payment_channel', current_app.config.get('PAYMENT_CHANNEL', 'alipay')) or 'alipay').strip().lower()


def _safe_days(value, default=7):
    try:
        days = int(value or default)
    except (TypeError, ValueError):
        days = default
    return max(1, min(days, 30))


def _build_export_file_response(article):
    title = (article.title or 'exported_article').strip() or 'exported_article'
    format_type = (article.format_type or 'txt').strip().lower()
    content = article.content or ''

    if format_type == 'txt':
        buffer = BytesIO()
        buffer.write(content.encode('utf-8'))
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name=f'{title}.txt', mimetype='text/plain')

    if format_type == 'docx':
        if not DOCX_AVAILABLE:
            return jsonify({'error': 'DOCX 格式暂不可用，请安装 python-docx'}), 400
        doc = Document()
        doc.add_heading(title, 0)
        for para in content.split('\n\n'):
            if para.strip():
                doc.add_paragraph(para.strip())
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f'{title}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    if format_type == 'pdf':
        if not PDF_AVAILABLE:
            return jsonify({'error': 'PDF 格式暂不可用，请安装 reportlab'}), 400
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        for para in content.split('\n\n'):
            if para.strip():
                story.append(Paragraph(para.strip(), styles['Normal']))
        doc.build(story)
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name=f'{title}.pdf', mimetype='application/pdf')

    buffer = BytesIO()
    buffer.write(content.encode('utf-8'))
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name=f'{title}.txt', mimetype='text/plain')


def _compact_pem_body(value):
    if not value:
        return ''
    normalized = str(value).replace('\\n', '\n').strip()
    lines = []
    for raw_line in normalized.splitlines():
        line = raw_line.strip()
        if not line or line.startswith('-----BEGIN ') or line.startswith('-----END '):
            continue
        lines.append(line)
    return ''.join(lines)


def _normalize_alipay_public_key(public_key):
    return _compact_pem_body(public_key)


def _normalize_alipay_private_key(private_key):
    from cryptography.hazmat.primitives import serialization

    body = _compact_pem_body(private_key)
    if not body:
        return ''

    pem_candidates = [
        f'-----BEGIN PRIVATE KEY-----\n{body}\n-----END PRIVATE KEY-----'.encode('utf-8'),
        f'-----BEGIN RSA PRIVATE KEY-----\n{body}\n-----END RSA PRIVATE KEY-----'.encode('utf-8'),
    ]

    loaded_key = None
    for candidate in pem_candidates:
        try:
            loaded_key = serialization.load_pem_private_key(candidate, password=None)
            break
        except Exception:
            continue

    if loaded_key is None:
        return body

    pkcs1_pem = loaded_key.private_bytes(
        encoding=serialization.Encoding.PEM,
        format=serialization.PrivateFormat.TraditionalOpenSSL,
        encryption_algorithm=serialization.NoEncryption(),
    ).decode('utf-8')
    return _compact_pem_body(pkcs1_pem)


def _alipay_sdk_available():
    try:
        from alipay.aop.api.AlipayClientConfig import AlipayClientConfig  # noqa: F401
        from alipay.aop.api.DefaultAlipayClient import DefaultAlipayClient  # noqa: F401
        from alipay.aop.api.domain.AlipayTradePagePayModel import AlipayTradePagePayModel  # noqa: F401
        from alipay.aop.api.request.AlipayTradePagePayRequest import AlipayTradePagePayRequest  # noqa: F401
        from alipay.aop.api.util.SignatureUtils import get_sign_content, verify_with_rsa  # noqa: F401
        return True
    except Exception:
        return False


def _alipay_ready():
    config = current_app.config
    required = [
        'ALIPAY_APP_ID',
        'ALIPAY_PRIVATE_KEY',
        'ALIPAY_PUBLIC_KEY',
        'ALIPAY_NOTIFY_URL',
        'ALIPAY_RETURN_URL',
    ]
    return _alipay_sdk_available() and all(config.get(key) for key in required)


def _build_alipay_client():
    from alipay.aop.api.AlipayClientConfig import AlipayClientConfig
    from alipay.aop.api.DefaultAlipayClient import DefaultAlipayClient

    client_config = AlipayClientConfig()
    client_config.server_url = current_app.config['ALIPAY_GATEWAY']
    client_config.app_id = current_app.config['ALIPAY_APP_ID']
    client_config.app_private_key = _normalize_alipay_private_key(current_app.config['ALIPAY_PRIVATE_KEY'])
    client_config.alipay_public_key = _normalize_alipay_public_key(current_app.config['ALIPAY_PUBLIC_KEY'])
    return DefaultAlipayClient(alipay_client_config=client_config, logger=current_app.logger)


def _build_alipay_page_url(order):
    from alipay.aop.api.domain.AlipayTradePagePayModel import AlipayTradePagePayModel
    from alipay.aop.api.request.AlipayTradePagePayRequest import AlipayTradePagePayRequest

    client = _build_alipay_client()
    model = AlipayTradePagePayModel()
    model.out_trade_no = order.order_no
    model.total_amount = str(order.amount_yuan)
    model.subject = order.subject
    model.product_code = 'FAST_INSTANT_TRADE_PAY'

    req = AlipayTradePagePayRequest(biz_model=model)
    req.notify_url = current_app.config['ALIPAY_NOTIFY_URL']
    req.return_url = current_app.config['ALIPAY_RETURN_URL']
    return client.page_execute(req, http_method='GET')


def _verify_notify_signature(form_data):
    from alipay.aop.api.util.SignatureUtils import get_sign_content, verify_with_rsa

    sign = (form_data or {}).get('sign')
    if not sign:
        return False

    sign_payload = {
        key: value
        for key, value in form_data.items()
        if key not in ('sign', 'sign_type') and value is not None and value != ''
    }
    sign_content = get_sign_content(sign_payload)
    charset = (form_data.get('charset') or 'utf-8').lower()
    try:
        message = sign_content.encode(charset)
    except Exception:
        message = sign_content.encode('utf-8')
    return bool(verify_with_rsa(_normalize_alipay_public_key(current_app.config['ALIPAY_PUBLIC_KEY']), message, sign))


@payment_bp.route('/wallet/summary')
@login_required
def wallet_summary():
    wallet = get_or_create_wallet(current_user.id)
    db.session.commit()

    recent_orders = RechargeOrder.query.filter_by(user_id=current_user.id).order_by(RechargeOrder.created_at.desc()).limit(10).all()
    recent_ledger = WalletLedger.query.filter_by(user_id=current_user.id).order_by(WalletLedger.created_at.desc()).limit(100).all()
    can_view_usage = bool(getattr(current_user, 'is_admin', False))
    recent_usage = []
    if can_view_usage:
        recent_usage = LLMUsageRecord.query.filter_by(user_id=current_user.id).order_by(LLMUsageRecord.created_at.desc()).limit(12).all()

    return jsonify({
        'wallet': {
            'available_coins': wallet.available_coins,
            'reserved_coins': wallet.reserved_coins,
            'total_recharged_coins': wallet.total_recharged_coins,
            'total_consumed_coins': wallet.total_consumed_coins,
        },
        'pricing': {
            'coins_per_1000_tokens': current_app.config['COINS_PER_1000_TOKENS'],
            'coins_per_yuan': current_app.config['COINS_PER_YUAN'],
            'review_audit_coins': int(current_app.config.get('REVIEW_AUDIT_COINS', '2') or 2),
            'min_recharge_amount_yuan': current_app.config['MIN_RECHARGE_AMOUNT_YUAN'],
            'payments_ready': _alipay_ready(),
            'payment_channel': _payment_channel(),
            'can_view_usage': can_view_usage,
        },
        'recent_orders': [
            {
                'order_no': item.order_no,
                'amount_yuan': str(item.amount_yuan),
                'coins': item.coins,
                'status': item.status,
                'created_at': item.created_at.isoformat() if item.created_at else '',
                'paid_at': item.paid_at.isoformat() if item.paid_at else '',
            }
            for item in recent_orders
        ],
        'recent_ledger': [
            {
                'change_type': item.change_type,
                'related_usage_id': item.related_usage_id,
                'related_order_no': item.related_order_no,
                'usage_endpoint': item.usage_record.endpoint if item.usage_record else '',
                'available_delta': item.available_delta,
                'reserved_delta': item.reserved_delta,
                'available_after': item.available_after,
                'reserved_after': item.reserved_after,
                'remark': item.remark or '',
                'created_at': item.created_at.isoformat() if item.created_at else '',
            }
            for item in recent_ledger
        ],
        'recent_usage': [
            {
                'call_id': item.call_id,
                'endpoint': item.endpoint,
                'model': item.model,
                'total_tokens': item.total_tokens,
                'coins_reserved': item.coins_reserved,
                'coins_charged': item.coins_charged,
                'status': item.status,
                'created_at': item.created_at.isoformat() if item.created_at else '',
            }
            for item in recent_usage
        ]
    })


@payment_bp.route('/payment/alipay/create', methods=['POST'])
@login_required
def create_alipay_order():
    if _payment_channel() != 'alipay':
        return jsonify({'error': '当前支付通道已关闭在线支付，请联系管理员。'}), 400

    if not _alipay_ready():
        return jsonify({'error': '支付宝支付配置未完成，请先填写 .env 中的支付参数并安装依赖。'}), 400

    data = request.get_json(silent=True) or request.form.to_dict()
    try:
        amount_yuan = Decimal(str(data.get('amount_yuan', '0'))).quantize(Decimal('0.01'))
    except (InvalidOperation, TypeError):
        return jsonify({'error': '充值金额格式不正确'}), 400

    min_amount = Decimal(str(current_app.config['MIN_RECHARGE_AMOUNT_YUAN']))
    if amount_yuan < min_amount:
        return jsonify({'error': f'单次充值金额不能低于 {min_amount} 元'}), 400

    subject_prefix = current_app.config['PAYMENT_SUBJECT_PREFIX']
    subject = f'{subject_prefix} - {current_user.username}'

    try:
        order = create_recharge_order(current_user.id, amount_yuan=amount_yuan, subject=subject)
        pay_url = _build_alipay_page_url(order)
        log_user_action(current_user.id, 'create_recharge_order', f'order_no={order.order_no}, amount={amount_yuan}')
        db.session.commit()
    except Exception as error:
        db.session.rollback()
        current_app.logger.exception('创建支付宝订单失败')
        return jsonify({'error': '创建支付订单失败，请稍后重试或联系管理员。'}), 500

    return jsonify({
        'order_no': order.order_no,
        'amount_yuan': str(order.amount_yuan),
        'coins': order.coins,
        'status': order.status,
        'pay_url': pay_url,
    })


@payment_bp.route('/payment/alipay/notify', methods=['POST'])
def alipay_notify():
    form_data = request.form.to_dict(flat=True)
    current_app.logger.info('收到支付宝异步通知: out_trade_no=%s trade_status=%s', form_data.get('out_trade_no'), form_data.get('trade_status'))

    if not _alipay_ready():
        current_app.logger.error('支付宝通知处理失败：支付参数未就绪')
        return 'failure'

    try:
        if not _verify_notify_signature(form_data):
            current_app.logger.error('支付宝通知验签失败: %s', form_data)
            return 'failure'

        order_no = form_data.get('out_trade_no', '').strip()
        trade_status = (form_data.get('trade_status') or '').strip()
        app_id = (form_data.get('app_id') or '').strip()
        if not order_no or not trade_status:
            current_app.logger.error('支付宝通知缺少关键字段: %s', form_data)
            return 'failure'

        configured_app_id = (current_app.config.get('ALIPAY_APP_ID') or '').strip()
        if app_id and configured_app_id and app_id != configured_app_id:
            current_app.logger.error('支付宝通知 app_id 不匹配: incoming=%s expected=%s', app_id, configured_app_id)
            return 'failure'

        order = RechargeOrder.query.filter_by(order_no=order_no, channel='alipay').with_for_update().first()
        if not order:
            current_app.logger.error('支付宝通知对应订单不存在: %s', order_no)
            db.session.rollback()
            return 'failure'

        total_amount_raw = (form_data.get('total_amount') or '').strip()
        if total_amount_raw:
            try:
                notify_amount = Decimal(total_amount_raw).quantize(Decimal('0.01'))
            except (InvalidOperation, TypeError):
                current_app.logger.error('支付宝通知金额格式非法: order=%s total_amount=%s', order_no, total_amount_raw)
                return 'failure'
            if notify_amount != Decimal(str(order.amount_yuan)).quantize(Decimal('0.01')):
                current_app.logger.error(
                    '支付宝通知金额不匹配: order=%s notify=%s expected=%s',
                    order_no,
                    notify_amount,
                    order.amount_yuan,
                )
                return 'failure'

        if trade_status in ('TRADE_SUCCESS', 'TRADE_FINISHED'):
            apply_recharge(
                order,
                alipay_trade_no=form_data.get('trade_no') or order.alipay_trade_no,
                buyer_logon_id=form_data.get('buyer_logon_id') or order.buyer_logon_id,
                notify_payload=json.dumps(form_data, ensure_ascii=False)
            )
            log_user_action(order.user_id, 'recharge_paid', f'order_no={order.order_no}, amount={order.amount_yuan}, coins={order.coins}')
        elif trade_status == 'TRADE_CLOSED':
            if order.status == 'pending':
                order.status = 'closed'
                order.notify_payload = json.dumps(form_data, ensure_ascii=False)
        else:
            current_app.logger.warning('支付宝通知状态忽略: %s order=%s', trade_status, order_no)

        db.session.commit()
        return 'success'
    except Exception:
        db.session.rollback()
        current_app.logger.exception('支付宝通知处理异常')
        return 'failure'


@payment_bp.route('/payment/dev/recharge/<order_no>', methods=['POST'])
@login_required
def dev_mark_recharge_paid(order_no):
    if not current_app.debug:
        return jsonify({'error': '该接口仅用于本地开发环境'}), 403

    order = RechargeOrder.query.filter_by(order_no=order_no, user_id=current_user.id).first()
    if not order:
        return jsonify({'error': '订单不存在'}), 404

    try:
        apply_recharge(
            order,
            alipay_trade_no=f'DEV-{order.order_no}',
            buyer_logon_id=current_user.username,
            notify_payload='development-manual-recharge'
        )
        log_user_action(current_user.id, 'recharge_paid', f'order_no={order.order_no}, amount={order.amount_yuan}, coins={order.coins}')
        db.session.commit()
    except Exception as error:
        db.session.rollback()
        return jsonify({'error': f'开发充值入账失败：{error}'}), 500

    return jsonify({'ok': True, 'order_no': order.order_no, 'status': order.status, 'coins': order.coins})


@payment_bp.route('/wallet/feature-charge', methods=['POST'])
@login_required
def wallet_feature_charge():
    data = request.get_json(silent=True) or request.form.to_dict()
    feature_code = str(data.get('feature_code') or '').strip().lower()
    request_id = str(data.get('request_id') or '').strip()

    feature_map = {
        'review_audit': {
            'coins': int(current_app.config.get('REVIEW_AUDIT_COINS', '2') or 2),
            'remark': '自动审核功能扣费 2 代币'
        }
    }

    feature = feature_map.get(feature_code)
    if not feature:
        return jsonify({'error': '不支持的功能扣费类型'}), 400

    if not request_id:
        return jsonify({'error': '缺少请求标识'}), 400

    idempotency_key = f'feature:{feature_code}:{request_id}'

    try:
        wallet, _ledger = charge_feature_coins(
            current_user.id,
            feature_code=feature_code,
            coins=feature['coins'],
            idempotency_key=idempotency_key,
            remark=feature['remark']
        )
        log_user_action(current_user.id, 'feature_charge', f'feature={feature_code}, coins={feature["coins"]}')
        db.session.commit()
    except InsufficientBalanceError as error:
        db.session.rollback()
        return jsonify({'error': str(error)}), 402
    except Exception as error:
        db.session.rollback()
        current_app.logger.exception('功能扣费失败: feature=%s user_id=%s', feature_code, current_user.id)
        return jsonify({'error': f'功能扣费失败：{error}'}), 500

    return jsonify({
        'ok': True,
        'feature_code': feature_code,
        'charged_coins': feature['coins'],
        'available_coins': wallet.available_coins,
        'reserved_coins': wallet.reserved_coins,
    })


@payment_bp.route('/wallet/exported-articles')
@login_required
def wallet_exported_articles():
    days = _safe_days(request.args.get('days', '7'), default=7)
    keyword = str(request.args.get('keyword', '') or '').strip()
    start_at = datetime.utcnow() - timedelta(days=days)

    query = ExportedArticle.query.filter(
        ExportedArticle.user_id == current_user.id,
        ExportedArticle.created_at >= start_at,
    )
    if keyword:
        from app.admin import _escape_ilike
        query = query.filter(ExportedArticle.title.ilike(f'%{_escape_ilike(keyword)}%'))

    items = query.order_by(ExportedArticle.created_at.desc()).limit(200).all()
    return jsonify({
        'days': days,
        'keyword': keyword,
        'items': [
            {
                'id': item.id,
                'title': item.title,
                'format_type': item.format_type,
                'content_length': item.content_length,
                'created_at': item.created_at.isoformat() if item.created_at else '',
                'download_name': f"{item.title}.{item.format_type or 'txt'}",
            }
            for item in items
        ]
    })


@payment_bp.route('/wallet/exported-articles/<int:article_id>/download')
@login_required
def wallet_exported_article_download(article_id):
    start_at = datetime.utcnow() - timedelta(days=7)
    article = ExportedArticle.query.filter(
        ExportedArticle.id == article_id,
        ExportedArticle.user_id == current_user.id,
        ExportedArticle.created_at >= start_at,
    ).first()
    if not article:
        return jsonify({'error': '未找到可下载记录（仅支持最近 7 天）'}), 404

    return _build_export_file_response(article)
