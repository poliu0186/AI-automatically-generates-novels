import re
import time
import hashlib
import threading
from io import BytesIO
from openai import OpenAI
from flask import Blueprint, request, Response, current_app, send_file, has_app_context, stream_with_context
from flask_login import login_required, current_user

from app.activity_logging import log_user_action
from app.billing import (
    InsufficientBalanceError,
    get_or_create_wallet,
    estimate_tokens_from_text,
    finalize_usage_charge,
    make_request_id,
    max_tokens_for_coins,
    release_usage_reservation,
    reserve_usage_charge,
)
from app.extensions import db
from app.models import ExportedArticle

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

api_bp = Blueprint('api', __name__)

_ROUTE_STATE_LOCK = threading.Lock()
_ROUTE_STATE = {}


def get_client(endpoint, api_key):
    return OpenAI(base_url=endpoint, api_key=api_key)


def _split_pool_values(raw):
    if not raw:
        return []
    parts = [item.strip() for item in str(raw).replace('\n', ',').split(',')]
    return [item for item in parts if item]


def _build_route_candidates(endpoint_key, api_key_name):
    suffix = endpoint_key.rsplit('_', 1)[-1]
    endpoint_pool_name = f'API_ENDPOINT_POOL_{suffix}'
    api_key_pool_name = f'API_KEY_POOL_{suffix}'

    endpoint_values = _split_pool_values(current_app.config.get(endpoint_pool_name, ''))
    api_key_values = _split_pool_values(current_app.config.get(api_key_pool_name, ''))

    base_endpoint = (current_app.config.get(endpoint_key) or '').strip()
    base_key = (current_app.config.get(api_key_name) or '').strip()

    if base_endpoint and base_endpoint not in endpoint_values:
        endpoint_values.insert(0, base_endpoint)
    if base_key and base_key not in api_key_values:
        api_key_values.insert(0, base_key)

    if not endpoint_values or not api_key_values:
        return []

    candidates = []
    for endpoint in endpoint_values:
        for api_key in api_key_values:
            route_hash = hashlib.sha1(f'{endpoint}|{api_key}'.encode('utf-8')).hexdigest()[:12]
            candidates.append(
                {
                    'id': route_hash,
                    'endpoint': endpoint,
                    'api_key': api_key,
                    'endpoint_key': endpoint_key,
                    'api_key_name': api_key_name,
                }
            )
    return candidates


def _get_route_state(state_key):
    state = _ROUTE_STATE.get(state_key)
    if not state:
        state = {'rr_index': 0, 'fail_counts': {}, 'cooldown_until': {}}
        _ROUTE_STATE[state_key] = state
    return state


def _select_route(endpoint_key, api_key_name):
    candidates = _build_route_candidates(endpoint_key, api_key_name)
    if not candidates:
        raise RuntimeError(f'No route candidates found for {endpoint_key}/{api_key_name}')

    state_key = f'{endpoint_key}:{api_key_name}'
    now_ts = time.time()

    with _ROUTE_STATE_LOCK:
        state = _get_route_state(state_key)
        available = [
            item
            for item in candidates
            if float(state['cooldown_until'].get(item['id'], 0) or 0) <= now_ts
        ]
        if not available:
            available = candidates

        idx = state['rr_index'] % len(available)
        selected = available[idx]
        state['rr_index'] = (state['rr_index'] + 1) % max(len(available), 1)

    return selected, candidates


def _mark_route_result(endpoint_key, api_key_name, route_id, ok):
    state_key = f'{endpoint_key}:{api_key_name}'
    fail_threshold = max(int(current_app.config.get('API_ROUTE_FAILURE_THRESHOLD', 2) or 2), 1)
    cooldown_seconds = max(int(current_app.config.get('API_ROUTE_COOLDOWN_SECONDS', 30) or 30), 1)

    with _ROUTE_STATE_LOCK:
        state = _get_route_state(state_key)
        if ok:
            state['fail_counts'][route_id] = 0
            state['cooldown_until'][route_id] = 0
            return

        fail_count = int(state['fail_counts'].get(route_id, 0) or 0) + 1
        state['fail_counts'][route_id] = fail_count
        if fail_count >= fail_threshold:
            state['cooldown_until'][route_id] = time.time() + cooldown_seconds
            state['fail_counts'][route_id] = 0


def _create_completion_with_fallback(*, candidates, selected_route, model_name, prompt, logger):
    ordered = [selected_route] + [item for item in candidates if item['id'] != selected_route['id']]
    last_error = None

    for route in ordered:
        client = get_client(route['endpoint'], route['api_key'])
        try:
            completion = client.chat.completions.create(
                model=model_name,
                messages=[{'role': 'user', 'content': prompt}],
                stream=True
            )
            _mark_route_result(route['endpoint_key'], route['api_key_name'], route['id'], True)
            return completion, route
        except Exception as error:
            _mark_route_result(route['endpoint_key'], route['api_key_name'], route['id'], False)
            logger.warning('LLM route failed, try next route: endpoint=%s route=%s error=%s', route['endpoint'], route['id'], error)
            last_error = error

    raise RuntimeError(f'All routes failed: {last_error}')


def _safe_export_title(raw_title):
    title = re.sub(r'\s+', ' ', str(raw_title or 'generated_novel')).strip()
    if not title:
        title = 'generated_novel'
    return title[:120]


def _extract_chunk_usage(chunk):
    usage = getattr(chunk, 'usage', None)
    if not usage:
        return None
    prompt_tokens = int(getattr(usage, 'prompt_tokens', 0) or 0)
    completion_tokens = int(getattr(usage, 'completion_tokens', 0) or 0)
    total_tokens = int(getattr(usage, 'total_tokens', 0) or (prompt_tokens + completion_tokens))
    return {
        'prompt_tokens': prompt_tokens,
        'completion_tokens': completion_tokens,
        'total_tokens': total_tokens,
    }


def _stream_with_billing(*, endpoint_name, endpoint_key, api_key_name, model_name):
    data = request.get_json(silent=True) or {}
    prompt = data.get('prompt', '')
    request_id = (data.get('request_id') or '').strip() or make_request_id()
    current_app.logger.debug('Received prompt for %s: %s', endpoint_name, prompt)

    estimated_input_tokens = estimate_tokens_from_text(prompt)
    estimated_output_tokens = int(current_app.config.get('DEFAULT_ESTIMATED_OUTPUT_TOKENS', '1200') or 1200)
    estimated_total_tokens = estimated_input_tokens + estimated_output_tokens

    selected_route, route_candidates = _select_route(endpoint_key, api_key_name)

    try:
        usage_record = reserve_usage_charge(
            current_user.id,
            request_id=request_id,
            endpoint=endpoint_name,
            provider=selected_route['endpoint'],
            model=model_name,
            estimated_tokens=estimated_total_tokens,
            remark=f'{endpoint_name} 预占：输入约 {estimated_input_tokens} + 输出约 {estimated_output_tokens} tokens'
        )
        db.session.commit()

        # Budget ceiling for this call: reserved for this call + available right after reservation.
        # This prevents one request from consuming more than the user's current total affordable coins.
        wallet_after_reserve = get_or_create_wallet(current_user.id)
        budget_coins = int(usage_record.coins_reserved or 0) + int(wallet_after_reserve.available_coins or 0)
        max_billable_tokens = max_tokens_for_coins(budget_coins)
    except InsufficientBalanceError as error:
        db.session.rollback()
        return Response(str(error), status=402, mimetype='text/plain')
    except Exception as error:
        db.session.rollback()
        current_app.logger.exception('预占代币失败: endpoint=%s', endpoint_name)
        return Response(f'预占代币失败: {error}', status=500, mimetype='text/plain')

    logger = current_app.logger

    def generate_stream():
        completion_text_parts = []
        provider_usage = None
        hit_budget_limit = False
        route_used = selected_route

        def over_budget(next_chunk_text='', next_chunk_usage=None):
            if max_billable_tokens <= 0:
                return True

            if next_chunk_usage and int(next_chunk_usage.get('total_tokens', 0) or 0) > max_billable_tokens:
                return True

            estimated_prompt = estimate_tokens_from_text(prompt)
            estimated_completion = estimate_tokens_from_text(''.join(completion_text_parts) + (next_chunk_text or ''))
            return (estimated_prompt + estimated_completion) > max_billable_tokens

        try:
            completion, route_used = _create_completion_with_fallback(
                candidates=route_candidates,
                selected_route=selected_route,
                model_name=model_name,
                prompt=prompt,
                logger=logger,
            )
            logger.debug('Stream created successfully for %s with route=%s', endpoint_name, route_used['id'])
            for chunk in completion:
                chunk_usage = _extract_chunk_usage(chunk)
                if chunk_usage:
                    provider_usage = chunk_usage

                chunk_text = None
                if getattr(chunk, 'choices', None):
                    delta = getattr(chunk.choices[0], 'delta', None)
                    chunk_text = getattr(delta, 'content', None)

                if over_budget(next_chunk_text=chunk_text or '', next_chunk_usage=chunk_usage):
                    hit_budget_limit = True
                    break

                if chunk_text:
                    completion_text_parts.append(chunk_text)
                    yield chunk_text

            if hit_budget_limit:
                yield '\n\n[系统提示] 已达到本次可支付上限，已停止继续生成。请先充值后重试。'

            if provider_usage:
                prompt_tokens = provider_usage['prompt_tokens']
                completion_tokens = provider_usage['completion_tokens']
                total_tokens = provider_usage['total_tokens']
                usage_source = 'provider'
            else:
                completion_text = ''.join(completion_text_parts)
                prompt_tokens = estimate_tokens_from_text(prompt)
                completion_tokens = estimate_tokens_from_text(completion_text)
                total_tokens = prompt_tokens + completion_tokens
                usage_source = 'estimated'

            finalize_usage_charge(
                usage_record.id,
                prompt_tokens=prompt_tokens,
                completion_tokens=completion_tokens,
                total_tokens=total_tokens,
                usage_source=usage_source,
            )
            log_user_action(
                current_user.id,
                'ai_generate_completed',
                f'endpoint={endpoint_name}, model={model_name}, route={route_used["id"]}, total_tokens={total_tokens}, request_id={request_id}'
            )
            db.session.commit()
        except GeneratorExit:
            # Client closed stream early. Release reservation safely and stop silently.
            try:
                if has_app_context():
                    db.session.rollback()
                    release_usage_reservation(usage_record.id, reason=f'{endpoint_name} 客户端中断，释放预占代币')
                    db.session.commit()
                else:
                    logger.warning('Stream closed after app context ended: usage_id=%s', usage_record.id)
            except Exception:
                if has_app_context():
                    db.session.rollback()
                logger.exception('客户端中断后释放预占代币失败: usage_id=%s', usage_record.id)
            raise
        except Exception as error:
            if has_app_context():
                db.session.rollback()
            try:
                release_usage_reservation(usage_record.id, reason=f'{endpoint_name} 调用失败，释放预占代币: {error}')
                if has_app_context():
                    db.session.commit()
            except Exception:
                if has_app_context():
                    db.session.rollback()
                logger.exception('释放预占代币失败: usage_id=%s', usage_record.id)
            logger.exception('Error in %s stream', endpoint_name)
            yield f'Error: {str(error)}'

    return Response(stream_with_context(generate_stream()), mimetype='text/plain')


@api_bp.route('/gen', methods=['POST'])
@login_required
def generate():
    return _stream_with_billing(
        endpoint_name='/gen',
        endpoint_key='API_ENDPOINT_1',
        api_key_name='API_KEY_1',
        model_name=(current_app.config.get('MODEL_NAME_1') or 'glm-4.5-air').strip(),
    )


@api_bp.route('/gen2', methods=['POST'])
@login_required
def generate2():
    return _stream_with_billing(
        endpoint_name='/gen2',
        endpoint_key='API_ENDPOINT_2',
        api_key_name='API_KEY_2',
        model_name=(current_app.config.get('MODEL_NAME_2') or 'glm-4.5-air').strip(),
    )


@api_bp.route('/download', methods=['POST'])
@login_required
def download_novel():
    data = request.json
    content = data.get('content', '')
    format_type = data.get('format', 'txt')
    title = _safe_export_title(data.get('title', 'generated_novel'))

    if not content:
        return {'error': 'No content provided'}, 400

    content = re.sub(r'<[^>]+>', '', content)
    content_length = len(content)

    export_record = ExportedArticle(
        user_id=current_user.id,
        title=title,
        format_type=str(format_type or 'txt').lower(),
        content=content,
        content_length=content_length,
    )
    db.session.add(export_record)

    if format_type == 'txt':
        buffer = BytesIO()
        buffer.write(content.encode('utf-8'))
        buffer.seek(0)
        log_user_action(current_user.id, 'novel_export_download', f'title={title}, format=txt, length={content_length}')
        db.session.commit()
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"{title}.txt",
            mimetype='text/plain'
        )

    elif format_type == 'docx':
        if not DOCX_AVAILABLE:
            db.session.rollback()
            return {'error': 'DOCX format not available. Please install python-docx'}, 400

        doc = Document()
        doc.add_heading(title, 0)
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        log_user_action(current_user.id, 'novel_export_download', f'title={title}, format=docx, length={content_length}')
        db.session.commit()
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"{title}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    elif format_type == 'pdf':
        if not PDF_AVAILABLE:
            db.session.rollback()
            return {'error': 'PDF format not available. Please install reportlab'}, 400

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                story.append(Paragraph(para.strip(), styles['Normal']))

        doc.build(story)
        buffer.seek(0)
        log_user_action(current_user.id, 'novel_export_download', f'title={title}, format=pdf, length={content_length}')
        db.session.commit()
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"{title}.pdf",
            mimetype='application/pdf'
        )

    else:
        db.session.rollback()
        return {'error': 'Unsupported format'}, 400
